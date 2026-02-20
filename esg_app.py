import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
import io
import re
import os
import shutil
from datetime import datetime, date
from urllib.parse import quote

st.set_page_config(page_title="ESG-Scoring", layout="wide")

class ESGScoringBot:
    def __init__(self):
        self.data = {}
        self.current_field = 0
        self.uploaded_files = []  # Liste für hochgeladene Dateien
        self.upload_folder = None  # Ordner für Dateien
        # Mappings für die Bedeutungen der Zahlenwerte
        self.value_mappings = {
            "korruptionsrisiko": {
                "1": "Erhöhtes Risiko",
                "2": "Leicht erhöhtes Risiko",
                "3": "Nein, keine Hinweise oder normales Risiko",
                "4": "Leicht verringertes Risiko",
                "5": "Geringes Risiko",
                "6": "Die Information liegt nicht vor."
            },
            "unternehmensfuehrung": {
                "1": "Besser als der Durchschnitt",
                "2": "Etwas über Durchschnitt",
                "3": "Durchschnitt",
                "4": "Etwas unter Durchschnitt",
                "5": "Unter Durchschnitt",
                "6": "Die Information liegt nicht vor."
            },
            "esgg_verknuepfung": {
                "1": "Ja",
                "2": "Nein, wird nicht benötigt",
                "3": "Nein, ist noch anzulegen"
            }
        }
        self.fields = [
            {"name": "vorname", "label": "Vorname", "question": "Vorname", "detailed_question": "Bitte gib Deinen Vornamen an."},
            {"name": "nachname", "label": "Nachname", "question": "Nachname", "detailed_question": "Bitte gib Deinen Nachnamen an."},
            {"name": "kd_esgg", "label": "KD bzw. KD ESGG", "question": "3-6 stellige Nummer", "detailed_question": "Für welche KD soll das ESG-Scoring angelegt werden?"},
            {"name": "kurzbezeichnung", "label": "Kurzbezeichnung", "question": "Gemäß Kreda", "detailed_question": "Wie lautet die Kurzbezeichnung des Unternehmens (gemäß Kreda)?"},
            {"name": "leit_kd", "label": "Leit-KD", "question": "Digi-Akte Nummer", "detailed_question": "Was ist die Leit-KD (Digi-Akte Nummer)?"},
            {"name": "fertigstellungstermin", "label": "Fertigstellung", "question": "TT.MM.JJJJ", "detailed_question": "Wann soll das ESG-Scoring fertiggestellt sein?"},
            {"name": "kurze_begruendung", "label": "Begründung", "question": "Frist-Anlass", "detailed_question": "Wie lautet die kurze Begründung für das Scoring (Frist-Anlass)?"},
            {"name": "mitarbeiter_anzahl", "label": "Mitarbeiter", "question": "Positive Zahl", "detailed_question": "Wie viele Mitarbeiter hat das Unternehmen?"},
            {"name": "umsatz_teur", "label": "Umsatz TEUR", "question": "Positive Zahl", "detailed_question": "Wie hoch ist der Umsatz des Unternehmens (in TEUR)?"},
            {"name": "bilanzsumme_teur", "label": "Bilanzsumme", "question": "Positive Zahl", "detailed_question": "Wie hoch ist die Bilanzsumme (in TEUR)?"},
            {"name": "sonstiges", "label": "ESG-relevante Dokumente", "question": "Dateien hochladen", "detailed_question": "Bitte lade hier ESG-relevante Dokumente hoch (z.B. Nachhaltigkeitsberichte, Zertifikate, Richtlinien, etc.). Dies ist optional."},
            {"name": "korruptionsrisiko", "label": "Korruptionsrisiko", "question": "1-6 wählen", "detailed_question": "Wie bewerten Sie das Korruptionsrisiko?\n\n1 = Erhöhtes Risiko\n2 = Leicht erhöhtes Risiko\n3 = Nein, keine Hinweise oder normales Risiko\n4 = Leicht verringertes Risiko\n5 = Geringes Risiko\n6 = Die Information liegt nicht vor."},
            {"name": "begruendung_korruptionsrisiko", "label": "Begründung Korruption", "question": "Erklären", "detailed_question": "Bitte begründen Sie Ihre Bewertung zum Korruptionsrisiko."},
            {"name": "unternehmensfuehrung", "label": "Unternehmensführung", "question": "1-6 wählen", "detailed_question": "Wie bewerten Sie die Unternehmensführung?\n\n1 = Besser als der Durchschnitt\n2 = Etwas über Durchschnitt\n3 = Durchschnitt\n4 = Etwas unter Durchschnitt\n5 = Unter Durchschnitt\n6 = Die Information liegt nicht vor."},
            {"name": "begruendung_unternehmensfuehrung", "label": "Begründung Führung", "question": "Erklären", "detailed_question": "Bitte begründen Sie Ihre Bewertung der Unternehmensführung."},
            {"name": "esgg_verknuepfung", "label": "ESGG-Verknüpfung", "question": "1-3 wählen", "detailed_question": "Ist eine ESG-Verknüpfung angelegt worden?\n\n1 = Ja\n2 = Nein, wird nicht benötigt\n3 = Nein, ist noch anzulegen"},
            {"name": "esgg_betroffene_kds", "label": "Betroffene KDs", "question": "3-6 stellige KD-Nummern (komma-getrennt)", "detailed_question": "Bitte ESG-Verknüpfung anlegen für folgende KDs"}
        ]
        self.word_content = None
        self.excel_content = None
    
    def validate(self, field_name, value):
        if not value or not value.strip(): 
            # Für "sonstiges" ist optional, auch wenn Dateien hochgeladen sind
            if field_name == "sonstiges":
                return True
            return field_name == "sonstiges"
        value = value.strip()
        
        if field_name in ["vorname", "nachname"]:
            # Nur Buchstaben, Bindestrich und Leerzeichen erlaubt
            return bool(re.match(r"^[a-zA-ZäöüßÄÖÜ\s\-]+$", value)) and len(value) >= 2
        elif field_name in ["kd_esgg", "leit_kd"]:
            return bool(re.match(r"^\d{3,6}$", value))
        elif field_name == "fertigstellungstermin":
            # Das Datum kommt bereits im Format TT.MM.JJJJ vom Kalender
            if not re.match(r"^\d{2}\.\d{2}\.\d{4}$", value):
                return False
            # Prüfe, ob es ein gültiges Datum ist und nicht in der Vergangenheit liegt
            try:
                selected_date = datetime.strptime(value, "%d.%m.%Y").date()
                return selected_date >= date.today()
            except ValueError:
                return False
        elif field_name == "mitarbeiter_anzahl":
            return value.isdigit() and int(value) > 0
        elif field_name in ["umsatz_teur", "bilanzsumme_teur"]:
            try: 
                return float(value) > 0
            except: 
                return False
        elif field_name == "korruptionsrisiko":
            return value in ["1","2","3","4","5","6"]
        elif field_name == "unternehmensfuehrung":
            return value in ["1","2","3","4","5","6"]
        elif field_name == "esgg_verknuepfung":
            return value in ["1","2","3"]
        elif field_name == "esgg_betroffene_kds":
            # Validiert komma- oder leerzeichen-getrennte KD-Nummern (3-6 Ziffern)
            kds = re.split(r'[,;\s]+', value.strip())
            for kd in kds:
                if kd and not re.match(r"^\d{3,6}$", kd):
                    return False
            return True
        return True
    
    def get_field_label(self, field_name):
        """Gibt aussagekräftige Labels für Feldnamen zurück"""
        labels = {
            "vorname": "Vorname",
            "nachname": "Nachname",
            "kd_esgg": "KD bzw. Kundennummer ESGG",
            "kurzbezeichnung": "Kurzbezeichnung",
            "leit_kd": "Leit-KD (Digi-Akte Nummer)",
            "fertigstellungstermin": "Fertigstellungstermin",
            "kurze_begruendung": "Begründung",
            "mitarbeiter_anzahl": "Mitarbeiteranzahl",
            "umsatz_teur": "Umsatz (TEUR)",
            "bilanzsumme_teur": "Bilanzsumme (TEUR)",
            "sonstiges": "ESG-relevante Dokumente",
            "korruptionsrisiko": "Korruptionsrisiko",
            "begruendung_korruptionsrisiko": "Begründung Korruptionsrisiko",
            "unternehmensfuehrung": "Unternehmensführung",
            "begruendung_unternehmensfuehrung": "Begründung Unternehmensführung",
            "esgg_verknuepfung": "ESG-Verknüpfung",
            "esgg_betroffene_kds": "Betroffene Kundennummern"
        }
        return labels.get(field_name, field_name.replace('_', ' ').title())
    
    def get_display_value(self, field_name, value):
        """Übersetzt Zahlenwerte in ihre Bedeutungen und formatiert Tausenderpunkte"""
        if field_name in self.value_mappings:
            return self.value_mappings[field_name].get(value, value)
        
        # Tausenderpunkte für numerische Felder
        if field_name in ["mitarbeiter_anzahl", "umsatz_teur", "bilanzsumme_teur"]:
            try:
                num = float(value)
                return f"{num:,.0f}".replace(",", ".")
            except (ValueError, TypeError):
                return value
        
        return value
    
    def save_uploaded_files(self, kd_esgg):
        """Erstellt Ordner für Dateien und speichert diese"""
        if not self.uploaded_files:
            return None
        
        # Ordner erstellen: uploads/KD_XXXX/
        self.upload_folder = f"uploads/KD_{kd_esgg}"
        os.makedirs(self.upload_folder, exist_ok=True)
        
        saved_files = []
        for uploaded_file in self.uploaded_files:
            file_path = os.path.join(self.upload_folder, uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            saved_files.append((uploaded_file.name, file_path))
        
        return saved_files
    
    def generate_email_link(self):
        """Generiert einen mailto-Link mit vorgedefinierten Daten"""
        vorname = self.data.get('vorname', '')
        kd_esgg = self.data.get('kd_esgg', '')
        kurzbezeichnung = self.data.get('kurzbezeichnung', '')
        
        to = "esg-scoring@ikb.de"
        subject = f"ESG-Scoring Auftrag für KD {kd_esgg}"
        
        body = f"""Liebes ESG-Scoring-Team,

anbei sende ich Euch den Scoring-Auftrag für KD {kd_esgg} ({kurzbezeichnung}) mit der Bitte um Bearbeitung.

Bitte meldet Euch, wenn Ihr Rückfragen habt.

Viele Grüße
{vorname}"""
        
        # URL-Encoding für mailto
        mailto_link = f"mailto:{to}?subject={quote(subject)}&body={quote(body)}"
        return mailto_link
    
    def create_documents(self):
        timestamp = datetime.now().strftime("%d%m%Y_%H%M")
        
        # E-Mail-Adresse generieren
        vorname = self.data.get('vorname', '').lower().strip()
        nachname = self.data.get('nachname', '').lower().strip()
        email = f"{vorname}.{nachname}@ikb.de" if vorname and nachname else "E-Mail nicht verfügbar"
        
        # KD-ESGG auslesen
        kd_esgg = self.data.get('kd_esgg', 'N/A')
        
        doc = Document()
        doc.add_heading(f"ESG-Auftrag für KD {kd_esgg}", 0)
        doc.add_paragraph(f"Auftraggeber: {email}")
        doc.add_paragraph(f"Erfassungsdatum: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        
        # Hochgeladene Dateien speichern und hinzufügen
        if self.uploaded_files:
            saved_files = self.save_uploaded_files(kd_esgg)
            if saved_files:
                doc.add_heading("Anhänge", level=2)
                for file_name, file_path in saved_files:
                    # Link zur Datei hinzufügen
                    doc.add_paragraph(f"📎 {file_name} → {file_path}", style='List Bullet')
        
        sections = {
            "KUNDENDATEN": ["kd_esgg", "kurzbezeichnung", "leit_kd", "fertigstellungstermin", "kurze_begruendung"],
            "UNTERNEHMENSDATEN": ["mitarbeiter_anzahl", "umsatz_teur", "bilanzsumme_teur", "sonstiges"],
            "RISIKOBEWERTUNG": ["korruptionsrisiko", "begruendung_korruptionsrisiko"],
            "CORPORATE GOVERNANCE": ["unternehmensfuehrung", "begruendung_unternehmensfuehrung"],
            "ESGG-VERKNUEPFUNG": ["esgg_verknuepfung", "esgg_betroffene_kds"]
        }
        
        for section, fields in sections.items():
            doc.add_heading(section, level=1)
            for field in fields:
                value = self.data.get(field, "—")
                display_value = self.get_display_value(field, value)
                field_label = self.get_field_label(field)
                doc.add_paragraph(f"{field_label}: {display_value}")
        
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        self.word_content = doc_io.getvalue()
        
        # Excel mit nur KD-ESGG und Kurzbezeichnung erstellen
        df = pd.DataFrame([{
            "KD-ESGG": self.data.get('kd_esgg', ''),
            "Kurzbezeichnung": self.data.get('kurzbezeichnung', '')
        }])
        excel_io = io.BytesIO()
        with pd.ExcelWriter(excel_io, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="ESG-Eingabe", index=False)
            # Spaltenbreiten anpassen
            worksheet = writer.sheets["ESG-Eingabe"]
            worksheet.column_dimensions['A'].width = 20
            worksheet.column_dimensions['B'].width = 40
        excel_io.seek(0)
        self.excel_content = excel_io.getvalue()

def main():
    st.title("🌿 ESG-Scoring Erfassung")
    
    if "bot" not in st.session_state:
        st.session_state.bot = ESGScoringBot()
    
    bot = st.session_state.bot
    
    # Progress Bar - nur während Eingabe anzeigen
    if bot.current_field < 17:
        progress = max(0.0, min(1.0, bot.current_field / 17.0))
        st.progress(progress)
        st.metric("Schritt", f"{bot.current_field + 1} von 17")
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Steuerung")
        if st.button("🔄 Neustart"):
            st.session_state.bot = ESGScoringBot()
            st.rerun()
    
    if bot.current_field < 17:
        field = bot.fields[bot.current_field]
        
        # Ausformulierte Frage anzeigen
        st.markdown(f"### ❓ {field['detailed_question']}")
        st.divider()
        
        # Spezielle Behandlung für Fertigstellungstermin mit Kalender
        if field['name'] == "fertigstellungstermin":
            selected_date = st.date_input(
                "📅 Fertigstellungsdatum wählen (TT.MM.JJJJ):",
                value=date.today(),
                min_value=date.today(),
                format="DD.MM.YYYY",
                key="date_picker"
            )
            user_input = selected_date.strftime("%d.%m.%Y") if selected_date else ""
            # Anzeige des gewählten Datums
            if user_input:
                st.info(f"✅ Gewähltes Datum: **{user_input}**")
        # Spezielle Behandlung für Sonstiges mit Datei-Upload
        elif field['name'] == "sonstiges":
            st.write("📁 Du kannst hier optional Dateien hochladen (alle Dateitypen erlaubt):")
            uploaded_files = st.file_uploader(
                "Dateien hochladen:",
                accept_multiple_files=True,
                key=f"upload_{bot.current_field}"
            )
            if uploaded_files:
                bot.uploaded_files = uploaded_files
                user_input = f"{len(uploaded_files)} Datei(en) hochgeladen"
                st.success(f"✅ {user_input}")
            else:
                user_input = ""
                st.info("Keine Dateien hochgeladen (optional)")
        else:
            user_input = st.text_input("Ihre Eingabe:", placeholder=field['question'], key=f"input_{bot.current_field}")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Speichern & Weiter", type="primary"):
                if bot.validate(field['name'], user_input):
                    bot.data[field['name']] = user_input
                    # Wenn Korruptionsrisiko = 3 oder Unternehmensführung = 3, nächste Frage (Begründung) überspringen
                    if (field['name'] == 'korruptionsrisiko' and user_input == '3') or (field['name'] == 'unternehmensfuehrung' and user_input == '3'):
                        bot.current_field += 2
                    # Wenn ESG-Verknüpfung = 1 oder 2, nächste Frage (betroffene KDs) überspringen
                    elif field['name'] == 'esgg_verknuepfung' and user_input in ['1', '2']:
                        bot.current_field += 2
                    else:
                        bot.current_field += 1
                    st.success("✓ Gespeichert!")
                    st.rerun()
                else:
                    st.error(f"❌ Format prüfen: {field['question']}")
        with col2:
            if st.button("↩️ Zurück") and bot.current_field > 0:
                bot.current_field -= 1
                st.rerun()
    else:
        st.success("🎉 Alle Daten erfasst!")
        
        # Session State für Bearbeitung initialisieren
        if "edit_field" not in st.session_state:
            st.session_state.edit_field = None
        
        # Wenn kein Feld in Bearbeitung ist, Übersicht anzeigen
        if st.session_state.edit_field is None:
            st.subheader("📋 Übersicht")
            for field in bot.fields:
                value = bot.data.get(field['name'], "—")
                display_value = bot.get_display_value(field['name'], value)
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.metric(field['label'], display_value)
                with col2:
                    if st.button("✏️ Bearbeiten", key=f"edit_{field['name']}"):
                        st.session_state.edit_field = field['name']
                        st.rerun()
            
            st.divider()
            if st.button("📄 Dokumente erstellen", type="primary"):
                bot.create_documents()
                st.balloons()
                st.success("Dokumente bereit!")
            
            col1, col2, col3 = st.columns(3)
            if bot.word_content:
                col1.download_button(
                    label="📝 Word",
                    data=bot.word_content,
                    file_name=f"ESG-Auftrag_{datetime.now().strftime('%d%m%Y_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            if bot.excel_content:
                col2.download_button(
                    label="📊 Excel",
                    data=bot.excel_content,
                    file_name=f"ESG-Eingabe_{datetime.now().strftime('%d%m%Y_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            
            # Email-Link
            if bot.word_content:
                mailto_link = bot.generate_email_link()
                col3.markdown(
                    f'<a href="{mailto_link}" target="_blank"><button style="width:100%; padding:10px; background-color:#4CAF50; color:white; border:none; border-radius:4px; cursor:pointer; font-weight:bold;">📧 E-Mail öffnen</button></a>',
                    unsafe_allow_html=True
                )
        else:
            # Feld wird bearbeitet
            field_name = st.session_state.edit_field
            field = next((f for f in bot.fields if f['name'] == field_name), None)
            
            if field:
                st.subheader(f"✏️ Bearbeite: {field['label']}")
                st.divider()
                
                # Ausformulierte Frage anzeigen
                st.markdown(f"### ❓ {field['detailed_question']}")
                st.divider()
                
                # Spezielle Behandlung für Fertigstellungstermin mit Kalender
                if field['name'] == "fertigstellungstermin":
                    current_value = bot.data.get(field['name'], date.today())
                    try:
                        current_date = datetime.strptime(current_value, "%d.%m.%Y").date()
                    except:
                        current_date = date.today()
                    
                    selected_date = st.date_input(
                        "📅 Fertigstellungsdatum wählen (TT.MM.JJJJ):",
                        value=current_date,
                        min_value=date.today(),
                        format="DD.MM.YYYY",
                        key="date_picker_edit"
                    )
                    user_input = selected_date.strftime("%d.%m.%Y") if selected_date else ""
                    if user_input:
                        st.info(f"✅ Gewähltes Datum: **{user_input}**")
                # Spezielle Behandlung für Sonstiges mit Datei-Upload
                elif field['name'] == "sonstiges":
                    st.write("📁 Du kannst hier optional Dateien hochladen (alle Dateitypen erlaubt):")
                    uploaded_files = st.file_uploader(
                        "Dateien hochladen:",
                        accept_multiple_files=True,
                        key=f"upload_edit_{bot.current_field}"
                    )
                    if uploaded_files:
                        bot.uploaded_files = uploaded_files
                        user_input = f"{len(uploaded_files)} Datei(en) hochgeladen"
                        st.success(f"✅ {user_input}")
                    else:
                        user_input = ""
                        st.info("Keine Dateien hochgeladen (optional)")
                else:
                    current_value = bot.data.get(field['name'], "")
                    user_input = st.text_input("Neue Eingabe:", value=current_value, placeholder=field['question'], key="edit_input")
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("✅ Speichern", type="primary"):
                        if bot.validate(field['name'], user_input):
                            bot.data[field['name']] = user_input
                            st.session_state.edit_field = None
                            st.success("✓ Änderungen gespeichert!")
                            st.rerun()
                        else:
                            st.error(f"❌ Format prüfen: {field['question']}")
                with col2:
                    if st.button("❌ Abbrechen"):
                        st.session_state.edit_field = None
                        st.rerun()

if __name__ == "__main__":
    main()
