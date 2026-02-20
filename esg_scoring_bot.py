"""
ESG-Scoring Bot - Verbesserte Version
Erfasst systematisch ESG-Daten mit Validierung und Export (Word/Excel)
"""

import streamlit as st
import pandas as pd
from docx import Document
import io
import re
from datetime import datetime
import logging

# Logging Setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

st.set_page_config(page_title="ESG-Scoring", layout="wide")


class ESGScoringBot:
    """Verwaltet ESG-Datenerfassung mit Validierung und Export"""
    
    def __init__(self):
        self.data = {}
        self.current_field = 0
        self.fields = [
            {"name": "kd_esgg", "label": "KD bzw. KD ESGG", "question": "3-6 stellige Nummer", "type": "number"},
            {"name": "kurzbezeichnung", "label": "Kurzbezeichnung", "question": "Gemäß Kreda", "type": "text"},
            {"name": "leit_kd", "label": "Leit-KD", "question": "Digi-Akte Nummer", "type": "number"},
            {"name": "fertigstellungstermin", "label": "Fertigstellung", "question": "TT.MM.JJJJ", "type": "date"},
            {"name": "kurze_begruendung", "label": "Begründung", "question": "Frist-Anlass", "type": "text"},
            {"name": "mitarbeiter_anzahl", "label": "Mitarbeiter", "question": "Positive Zahl", "type": "number_positive"},
            {"name": "umsatz_teur", "label": "Umsatz TEUR", "question": "Positive Zahl", "type": "float_positive"},
            {"name": "bilanzsumme_teur", "label": "Bilanzsumme", "question": "Positive Zahl", "type": "float_positive"},
            {"name": "sonstiges", "label": "Sonstiges", "question": "Optional", "type": "text_optional"},
            {"name": "korruptionsrisiko", "label": "Korruptionsrisiko", "question": "1-6 wählen", "type": "rating"},
            {"name": "begruendung_korruptionsrisiko", "label": "Begründung Korruption", "question": "Erklären", "type": "text"},
            {"name": "unternehmensfuehrung", "label": "Unternehmensführung", "question": "1-6 wählen", "type": "rating"},
            {"name": "begruendung_unternehmensfuehrung", "label": "Begründung Führung", "question": "Erklären", "type": "text"},
            {"name": "esgg_verknuepfung", "label": "ESGG-Verknüpfung", "question": "1-3 wählen", "type": "rating_3"},
            {"name": "esgg_betroffene_kds", "label": "Betroffene KDs", "question": "3-6 stellige KD-Nummern (komma-getrennt)", "type": "kd_list"}
        ]
        self.word_content = None
        self.excel_content = None
        self.validation_errors = []
    
    def validate(self, field_name: str, value: str) -> bool:
        """
        Validiert die Eingabe basierend auf Feldtyp
        
        Args:
            field_name: Name des Feldes
            value: Eingabewert
        
        Returns:
            bool: True wenn valide, sonst False
        """
        try:
            self.validation_errors = []
            
            if not value or not value.strip():
                # Nur "sonstiges" und Text-Optional sind optional
                field_type = next((f["type"] for f in self.fields if f["name"] == field_name), None)
                return field_type in ["text_optional", "sonstiges"]
            
            value = value.strip()
            
            # Validierung nach Feldtyp
            if field_name in ["kd_esgg", "leit_kd"]:
                if not re.match(r"^\d{3,6}$", value):
                    self.validation_errors.append(f"Muss 3-6 Ziffern sein, nicht: {value}")
                    return False
            
            elif field_name == "fertigstellungstermin":
                if not re.match(r"^\d{2}\.\d{2}\.\d{4}$", value):
                    self.validation_errors.append("Format: TT.MM.JJJJ (z.B. 19.02.2026)")
                    return False
                # Zusätzliche Validierung: Gültiges Datum?
                try:
                    datetime.strptime(value, "%d.%m.%Y")
                except ValueError:
                    self.validation_errors.append("Ungültiges Datum")
                    return False
            
            elif field_name == "mitarbeiter_anzahl":
                if not value.isdigit() or int(value) <= 0:
                    self.validation_errors.append("Muss positive Zahl sein")
                    return False
            
            elif field_name in ["umsatz_teur", "bilanzsumme_teur"]:
                try:
                    val = float(value.replace(",", "."))
                    if val <= 0:
                        self.validation_errors.append("Muss positive Zahl sein")
                        return False
                except ValueError:
                    self.validation_errors.append("Keine gültige Zahl")
                    return False
            
            elif field_name in ["korruptionsrisiko", "unternehmensfuehrung"]:
                if value not in ["1", "2", "3", "4", "5", "6"]:
                    self.validation_errors.append("Muss zwischen 1-6 wählen")
                    return False
            
            elif field_name == "esgg_verknuepfung":
                if value not in ["1", "2", "3"]:
                    self.validation_errors.append("Muss zwischen 1-3 wählen")
                    return False
            
            elif field_name == "esgg_betroffene_kds":
                # Validiert komma- oder leerzeichen-getrennte KD-Nummern (3-6 Ziffern)
                kds = re.split(r'[,;\s]+', value.strip())
                for kd in kds:
                    if kd and not re.match(r"^\d{3,6}$", kd):
                        self.validation_errors.append(f"Ungültige KD-Nummer: {kd} (muss 3-6 Ziffern sein)")
                        return False
            
            return True
        
        except Exception as e:
            logger.error(f"Validierungsfehler für {field_name}: {e}")
            self.validation_errors = [f"Technischer Fehler: {str(e)}"]
            return False
    
    def create_documents(self) -> bool:
        """Erstellt Word- und Excel-Dokumente"""
        try:
            timestamp = datetime.now().strftime("%d%m%Y_%H%M")
            
            # === WORD DOCUMENT ===
            doc = Document()
            doc.add_heading("ESG-Scoring Auftrag", 0)
            doc.add_paragraph(f"Erfassungsdatum: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            doc.add_paragraph("")  # Leerzeile
            
            sections = {
                "KUNDENDATEN": ["kd_esgg", "kurzbezeichnung", "leit_kd", "fertigstellungstermin", "kurze_begruendung"],
                "UNTERNEHMENSDATEN": ["mitarbeiter_anzahl", "umsatz_teur", "bilanzsumme_teur", "sonstiges"],
                "RISIKOBEWERTUNG": ["korruptionsrisiko", "begruendung_korruptionsrisiko"],
                "CORPORATE GOVERNANCE": ["unternehmensfuehrung", "begruendung_unternehmensfuehrung"],
                "ESSG-VERKNUEPFUNG": ["essg_verknuepfung", "essg_betroffene_kds"]
            }
            
            for section, field_list in sections.items():
                doc.add_heading(section, level=1)
                for field_name in field_list:
                    # Finde Label für Feldname
                    field_label = next((f["label"] for f in self.fields if f["name"] == field_name), field_name)
                    value = self.data.get(field_name, "—")
                    doc.add_paragraph(f"{field_label}: {value}")
            
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            self.word_content = doc_io.getvalue()
            
            # === EXCEL SPREADSHEET ===
            df = pd.DataFrame(list(self.data.items()), columns=["Feld", "Wert"])
            excel_io = io.BytesIO()
            with pd.ExcelWriter(excel_io, engine="openpyxl") as writer:
                df.to_excel(writer, sheet_name="ESG-Eingabe", index=False)
            excel_io.seek(0)
            self.excel_content = excel_io.getvalue()
            
            logger.info("Dokumente erfolgreich erstellt")
            return True
        
        except Exception as e:
            logger.error(f"Fehler beim Erstellen von Dokumenten: {e}")
            return False


def main():
    """Hauptfunktion für Streamlit App"""
    st.title("🌿 ESG-Scoring Erfassung")
    st.markdown("Erfassen Sie systematisch ESG-Daten. Alle Felder müssen validiert sein.")
    
    # Session State initialisieren
    if "bot" not in st.session_state:
        st.session_state.bot = ESGScoringBot()
    
    bot = st.session_state.bot
    
    # Progress Bar
    progress = max(0.0, min(1.0, bot.current_field / len(bot.fields)))
    st.progress(progress)
    col1, col2 = st.columns([3, 1])
    col1.metric("Fortschritt", f"{bot.current_field} / {len(bot.fields)}")
    col2.metric("Status", "✅ Aktiv" if bot.current_field < len(bot.fields) else "🎉 Fertig")
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Steuerung")
        if st.button("🔄 Neustart", use_container_width=True):
            st.session_state.bot = ESGScoringBot()
            st.rerun()
        
        st.divider()
        st.info(f"📊 Bearbeitete Felder: {len(bot.data)}")
    
    # Hauptbereich
    if bot.current_field < len(bot.fields):
        field = bot.fields[bot.current_field]
        
        st.markdown(f"### 📝 **{field['label']}**")
        st.caption(f"Hinweis: {field['question']}")
        
        user_input = st.text_input(
            "Ihre Eingabe:",
            key=f"input_{bot.current_field}",
            placeholder=field['question']
        )
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("✅ Speichern & Weiter", type="primary", use_container_width=True):
                if bot.validate(field['name'], user_input):
                    bot.data[field['name']] = user_input.strip()
                    bot.current_field += 1
                    st.success("✓ Gespeichert!", icon="✅")
                    st.rerun()
                else:
                    error_msg = " | ".join(bot.validation_errors) if bot.validation_errors else "Ungültiges Format"
                    st.error(f"❌ {error_msg}")
        
        with col2:
            if st.button("↩️ Zurück", use_container_width=True) and bot.current_field > 0:
                bot.current_field -= 1
                st.rerun()
        
        with col3:
            if st.button("🏠 Übersicht", use_container_width=True):
                st.session_state.show_preview = not st.session_state.get("show_preview", False)
                st.rerun()
        
        # Daten-Übersicht falls gewünscht
        if st.session_state.get("show_preview", False):
            st.divider()
            st.subheader("📋 Bisherige Eingaben")
            for field_info in bot.fields[:bot.current_field]:
                value = bot.data.get(field_info['name'], "—")
                st.caption(f"**{field_info['label']}:** {value}")
    
    else:
        # Abschlussseite
        st.success("🎉 Alle Daten erfasst!", icon="✅")
        
        st.subheader("📋 Finale Übersicht")
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### Kundendaten")
            for field in bot.fields[:5]:
                st.text(f"{field['label']}: {bot.data.get(field['name'], '—')}")
        
        with col2:
            st.markdown("#### Bewertungen")
            for field in bot.fields[9:]:
                st.text(f"{field['label']}: {bot.data.get(field['name'], '—')}")
        
        st.divider()
        
        # Dokumente erstellen
        if st.button("📄 Dokumente erstellen", type="primary", use_container_width=True):
            with st.spinner("Erstelle Dokumente..."):
                if bot.create_documents():
                    st.balloons()
                    st.success("Dokumente erfolgreich erstellt!")
        
        # Download Buttons
        if bot.word_content and bot.excel_content:
            st.divider()
            col1, col2 = st.columns(2)
            
            with col1:
                st.download_button(
                    label="📝 Word-Dokument",
                    data=bot.word_content,
                    file_name=f"ESG-Auftrag_{datetime.now().strftime('%d%m%Y_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col2:
                st.download_button(
                    label="📊 Excel-Datei",
                    data=bot.excel_content,
                    file_name=f"ESG-Eingabe_{datetime.now().strftime('%d%m%Y_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )


if __name__ == "__main__":
    main()
